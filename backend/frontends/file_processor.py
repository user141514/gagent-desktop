import hashlib
import mimetypes
import os
import re
from pathlib import Path


PROJECT_ROOT = Path(__file__).resolve().parents[1]
UPLOAD_ROOT = PROJECT_ROOT / "temp" / "uploads"

TEXT_SUFFIXES = {
    ".txt",
    ".md",
    ".py",
    ".json",
    ".csv",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".log",
    ".sql",
    ".js",
    ".ts",
    ".html",
    ".css",
    ".xml",
}
SUPPORTED_UPLOAD_SUFFIXES = sorted(TEXT_SUFFIXES | {".pdf", ".docx"})
PROMPT_WARN_CHAR_THRESHOLD = 50_000
ATTACHMENT_CONTEXT_START = "<uploaded_file_context>"
ATTACHMENT_CONTEXT_END = "</uploaded_file_context>"

_SENSITIVE_PATTERNS = [
    re.compile(r"(^|[\\/])\.(env|envrc)(\.|$)", re.IGNORECASE),
    re.compile(r"\.(pem|key|p12|pfx|cert|crt|der|p8)$", re.IGNORECASE),
    re.compile(r"(credential|secret|passwd|password|token|private_key)", re.IGNORECASE),
    re.compile(r"(id_rsa|id_dsa|id_ecdsa|id_ed25519)(\.pub)?$", re.IGNORECASE),
    re.compile(r"(\.netrc|\.pgpass|\.htpasswd)$", re.IGNORECASE),
]


def build_upload_id(filename, data):
    h = hashlib.sha1()
    h.update((filename or "").encode("utf-8", errors="ignore"))
    h.update(b"\0")
    h.update(data or b"")
    return h.hexdigest()[:16]


def guess_kind(filename, mime_type=""):
    suffix = Path(filename or "").suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return "text"
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if (mime_type or "").startswith("text/"):
        return "text"
    return "unsupported"


def format_size(num_bytes):
    if num_bytes < 1024:
        return f"{num_bytes} B"
    if num_bytes < 1024 * 1024:
        return f"{num_bytes / 1024:.1f} KB"
    return f"{num_bytes / (1024 * 1024):.1f} MB"


def _is_sensitive_name(filename):
    return any(p.search(filename or "") for p in _SENSITIVE_PATTERNS)


def _safe_filename(name):
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", name or "upload")
    return cleaned[:120] or "upload"


def _decode_text_bytes(data):
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _normalize_text(text):
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _compress_text(text, max_chars=12000):
    text = _normalize_text(text)
    if len(text) <= max_chars:
        return text

    lines = [line.rstrip() for line in text.splitlines()]
    non_empty = [line for line in lines if line.strip()]
    headings = [
        line for line in non_empty
        if len(line) <= 120 and (
            line.startswith("#")
            or line.endswith(":")
            or re.match(r"^(\d+(\.\d+)*|[A-Z])[\). ]", line)
        )
    ][:20]

    head_chars = max_chars // 3
    tail_chars = max_chars // 5
    middle_budget = max_chars - head_chars - tail_chars - 256

    head = text[:head_chars].strip()
    tail = text[-tail_chars:].strip()
    middle = "\n".join(headings)
    if len(middle) > middle_budget:
        middle = middle[:middle_budget].strip()

    parts = [head]
    if middle:
        parts.append("[Key sections]\n" + middle)
    parts.append("[Tail excerpt]\n" + tail)
    return "\n\n".join(part for part in parts if part).strip()


def _persist_upload(filename, data, file_id):
    UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)
    target_dir = UPLOAD_ROOT / file_id
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path = target_dir / _safe_filename(filename)
    target_path.write_bytes(data)
    return str(target_path)


def extract_text_file(path):
    data = Path(path).read_bytes()
    text = _decode_text_bytes(data)
    text = _normalize_text(text)
    return {
        "raw_text": text,
        "distilled_text": _compress_text(text, max_chars=12000),
        "preview_text": text[:1600] if text else "",
    }


def extract_docx_file(path):
    from docx import Document

    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        if style.startswith("Heading 1"):
            lines.append(f"# {text}")
        elif style.startswith("Heading 2"):
            lines.append(f"## {text}")
        elif style.startswith("Heading 3"):
            lines.append(f"### {text}")
        elif style.startswith("List"):
            lines.append(f"- {text}")
        else:
            lines.append(text)

    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        header = "| " + " | ".join(rows[0]) + " |"
        sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
        lines.extend([header, sep])
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")

    text = "\n".join(lines)
    text = _normalize_text(text)
    return {
        "raw_text": text,
        "distilled_text": _compress_text(text, max_chars=12000),
        "preview_text": text[:1600] if text else "",
    }


def extract_pdf_file(path):
    import fitz

    doc = fitz.open(path)
    total_pages = doc.page_count
    toc = doc.get_toc(simple=True) or []
    selected_pages = list(range(total_pages))
    if total_pages > 40:
        selected_pages = sorted(
            set(list(range(min(5, total_pages))) + list(range(max(0, total_pages - 3), total_pages)))
        )
        if toc:
            for _, _, page_no in toc[:12]:
                idx = max(0, min(total_pages - 1, int(page_no) - 1))
                selected_pages.append(idx)
        step = max(1, total_pages // 12)
        selected_pages.extend(range(0, total_pages, step))
        selected_pages = sorted(set(selected_pages))[:24]

    chunks = []
    for page_idx in selected_pages:
        page = doc.load_page(page_idx)
        page_text = _normalize_text(page.get_text("text"))
        if page_text:
            chunks.append(f"[Page {page_idx + 1}]\n{page_text}")
    full_text = "\n\n".join(chunks).strip()

    toc_text = ""
    if toc:
        toc_lines = [f"{'  ' * max(0, level - 1)}- {title} (p.{page_no})" for level, title, page_no in toc[:40]]
        toc_text = "[Table of Contents]\n" + "\n".join(toc_lines)

    distilled_body = _compress_text(full_text, max_chars=10000)
    distilled_text = (
        f"PDF pages: {total_pages}\n"
        + (toc_text + "\n\n" if toc_text else "")
        + distilled_body
    ).strip()

    return {
        "raw_text": full_text,
        "distilled_text": distilled_text,
        "preview_text": (toc_text + "\n\n" + full_text[:1200]).strip(),
        "page_count": total_pages,
        "toc_count": len(toc),
    }


def process_uploaded_file(uploaded_file):
    filename = getattr(uploaded_file, "name", "upload")
    data = uploaded_file.getvalue()
    size = len(data)
    mime_type = getattr(uploaded_file, "type", "") or mimetypes.guess_type(filename)[0] or ""
    file_id = build_upload_id(filename, data)
    stored_path = _persist_upload(filename, data, file_id)
    kind = guess_kind(filename, mime_type)

    meta = {
        "id": file_id,
        "name": filename,
        "size": size,
        "size_label": format_size(size),
        "mime": mime_type,
        "kind": kind,
        "stored_path": stored_path,
        "status": "ready",
        "preview_text": "",
        "distilled_text": "",
        "raw_text_length": 0,
        "warning": "",
    }

    warnings = []
    if _is_sensitive_name(filename):
        warnings.append("文件名看起来像敏感配置或密钥文件，发送给模型前请再次确认。")

    if kind == "unsupported":
        meta["status"] = "error"
        meta["warning"] = "暂不支持该文件类型，当前仅支持文本、PDF、DOCX。"
        return meta

    try:
        if kind == "text":
            result = extract_text_file(stored_path)
        elif kind == "pdf":
            result = extract_pdf_file(stored_path)
        else:
            result = extract_docx_file(stored_path)
        meta["preview_text"] = result.get("preview_text", "")
        meta["distilled_text"] = result.get("distilled_text", "")
        meta["raw_text_length"] = len(result.get("raw_text", ""))
        if meta["raw_text_length"] > PROMPT_WARN_CHAR_THRESHOLD:
            warnings.append("文件内容较大，已自动压缩后再并入当前提问。")
        if kind == "pdf" and result.get("page_count", 0) > 100:
            warnings.append(f"PDF 共 {result['page_count']} 页，已做保守抽样压缩。")
    except Exception as e:
        meta["status"] = "error"
        meta["warning"] = f"处理失败: {e}"
        return meta

    if warnings:
        meta["warning"] = "\n".join(warnings)
    return meta


def build_attachment_prompt(attachments, max_total_chars=18000):
    ready = [a for a in attachments if a.get("status") == "ready" and a.get("distilled_text")]
    if not ready:
        return ""

    per_file_budget = max(2000, max_total_chars // max(len(ready), 1))
    sections = []
    for meta in ready:
        body = meta["distilled_text"][:per_file_budget].strip()
        header = f"[Attachment] {meta['name']} | type={meta['kind']} | size={meta['size_label']}"
        sections.append(header + "\n" + body)

    joined = "\n\n".join(sections)
    if len(joined) > max_total_chars:
        joined = joined[:max_total_chars].rstrip()

    body = (
        "### Uploaded File Context\n"
        "The user uploaded the following files. Treat them as reference material for the current task. "
        "Prefer citing filenames when you rely on them. This block is valid only for the current request; "
        "later uploaded file context overrides earlier uploaded file context.\n\n"
        f"{joined}"
    ).strip()
    return f"{ATTACHMENT_CONTEXT_START}\n{body}\n{ATTACHMENT_CONTEXT_END}"


def strip_attachment_prompt(text):
    """Remove uploaded-file context blocks from text persisted in chat history."""
    if not isinstance(text, str) or not text:
        return text

    tagged = re.compile(
        rf"\n*{re.escape(ATTACHMENT_CONTEXT_START)}\s*.*?\s*{re.escape(ATTACHMENT_CONTEXT_END)}",
        re.DOTALL,
    )
    cleaned = tagged.sub("", text)

    # Backward compatibility for histories created before explicit tags existed.
    legacy_marker = "### Uploaded File Context"
    while legacy_marker in cleaned:
        start = cleaned.find(legacy_marker)
        suffix_start = cleaned.find("\n=== ASSISTANT ===", start)
        if suffix_start >= 0:
            cleaned = cleaned[:start].rstrip() + "\n" + cleaned[suffix_start:].lstrip()
        else:
            cleaned = cleaned[:start].rstrip()

    return cleaned.strip()
